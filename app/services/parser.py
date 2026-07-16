"""
Resume text extraction.

Pulls raw text out of an uploaded PDF or DOCX file so the analyzer can
run keyword and heuristic scoring against it.
"""

from __future__ import annotations

import io

import pdfplumber
from docx import Document

SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx"}


class UnsupportedFileType(ValueError):
    """Raised when the uploaded file isn't a resume format we can parse."""


def get_extension(filename: str) -> str:
    if "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Route to the right extractor based on file extension.

    Raises UnsupportedFileType if the extension isn't PDF/DOC/DOCX.
    """
    extension = get_extension(filename)

    if extension == ".pdf":
        return _extract_pdf_text(file_bytes)
    if extension in (".doc", ".docx"):
        return _extract_docx_text(file_bytes)

    raise UnsupportedFileType(
        f"Unsupported file type '{extension or 'unknown'}'. "
        f"Please upload a PDF or DOCX resume."
    )


def _extract_pdf_text(file_bytes: bytes) -> str:
    text_chunks: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_chunks.append(page_text)
    return "\n".join(text_chunks).strip()


def _extract_docx_text(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs).strip()